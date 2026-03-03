"""
文档管理器
"""
import asyncio
import hashlib
import os
import uuid
from datetime import datetime
from typing import Dict, Any, List, Optional
from pathlib import Path
import aiofiles
import PyPDF2
from docx import Document
from bs4 import BeautifulSoup
import markdown

from src.storage.database import get_db_session
from src.storage.models import KnowledgeDocument


class DocumentManager:
    """文档管理器"""
    
    def __init__(self, upload_dir: str = "uploads/documents"):
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        self.supported_formats = ['pdf', 'docx', 'txt', 'md', 'html']
    
    def validate_format(self, filename: str) -> bool:
        """验证文件格式"""
        file_extension = filename.split('.')[-1].lower()
        return file_extension in self.supported_formats
    
    async def upload_document(self, file_content: bytes, filename: str, category: str, tags: List[str] = None) -> Dict[str, Any]:
        """上传文档"""
        if tags is None:
            tags = []
        
        # 1. 验证文件格式
        if not self.validate_format(filename):
            raise ValueError(f"Unsupported file format: {filename}")
        
        # 2. 生成文件路径
        file_id = str(uuid.uuid4())
        file_extension = filename.split('.')[-1].lower()
        file_path = self.upload_dir / f"{file_id}.{file_extension}"
        
        # 3. 保存文件
        await self.save_file(file_content, file_path)
        
        # 4. 计算文件哈希
        file_hash = await self.calculate_file_hash(file_path)
        
        # 5. 解析文档内容
        content = await self.parse_document(file_path)
        
        # 6. 创建文档记录
        async for session in get_db_session():
            doc_record = KnowledgeDocument(
                document_id=file_id,
                filename=filename,
                file_path=str(file_path),
                file_size=len(file_content),
                file_hash=file_hash,
                category=category,
                tags=tags,
                content=content,
                status="processing"
            )
            
            session.add(doc_record)
            await session.commit()
            await session.refresh(doc_record)
            
            # 7. 触发向量化处理（由外部任务队列处理）
            # await self.trigger_vectorization(doc_record.id)
            
            return {
                "document_id": file_id,
                "filename": filename,
                "category": category,
                "tags": tags,
                "status": "processing",
                "created_at": doc_record.created_at
            }
    
    async def save_file(self, file_content: bytes, file_path: Path) -> None:
        """保存文件"""
        async with aiofiles.open(file_path, 'wb') as f:
            await f.write(file_content)
    
    async def calculate_file_hash(self, file_path: Path) -> str:
        """计算文件哈希"""
        hash_md5 = hashlib.md5()
        async with aiofiles.open(file_path, 'rb') as f:
            while chunk := await f.read(8192):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()
    
    async def parse_document(self, file_path: Path) -> str:
        """解析文档内容"""
        file_extension = file_path.suffix.lower().lstrip('.')
        
        parsers = {
            'pdf': self.parse_pdf,
            'docx': self.parse_docx,
            'txt': self.parse_txt,
            'md': self.parse_markdown,
            'html': self.parse_html
        }
        
        parser = parsers.get(file_extension)
        if not parser:
            raise ValueError(f"No parser for {file_extension}")
        
        return await parser(file_path)
    
    async def parse_pdf(self, file_path: Path) -> str:
        """解析PDF文件"""
        content = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                content += page.extract_text() + "\n"
        return content.strip()
    
    async def parse_docx(self, file_path: Path) -> str:
        """解析Word文档"""
        doc = Document(file_path)
        content = ""
        for paragraph in doc.paragraphs:
            content += paragraph.text + "\n"
        return content.strip()
    
    async def parse_txt(self, file_path: Path) -> str:
        """解析文本文件"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
            content = await f.read()
        return content
    
    async def parse_markdown(self, file_path: Path) -> str:
        """解析Markdown文件"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
            md_content = await f.read()
        # 转换为HTML然后提取文本
        html_content = markdown.markdown(md_content)
        soup = BeautifulSoup(html_content, 'html.parser')
        return soup.get_text()
    
    async def parse_html(self, file_path: Path) -> str:
        """解析HTML文件"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
            html_content = await f.read()
        soup = BeautifulSoup(html_content, 'html.parser')
        return soup.get_text()
    
    async def get_document(self, document_id: str) -> Optional[Dict[str, Any]]:
        """获取文档信息"""
        async for session in get_db_session():
            doc = await session.get(KnowledgeDocument, document_id)
            if doc:
                return {
                    "document_id": doc.document_id,
                    "filename": doc.filename,
                    "file_path": doc.file_path,
                    "file_size": doc.file_size,
                    "file_hash": doc.file_hash,
                    "category": doc.category,
                    "tags": doc.tags,
                    "content": doc.content,
                    "status": doc.status,
                    "created_at": doc.created_at,
                    "updated_at": doc.updated_at
                }
            return None
    
    async def update_document_status(self, document_id: str, status: str) -> bool:
        """更新文档状态"""
        async for session in get_db_session():
            doc = await session.get(KnowledgeDocument, document_id)
            if doc:
                doc.status = status
                doc.updated_at = datetime.utcnow()
                await session.commit()
                return True
            return False
    
    async def delete_document(self, document_id: str) -> bool:
        """删除文档"""
        async for session in get_db_session():
            doc = await session.get(KnowledgeDocument, document_id)
            if doc:
                # 删除文件
                try:
                    file_path = Path(doc.file_path)
                    if file_path.exists():
                        file_path.unlink()
                except Exception as e:
                    print(f"Warning: Failed to delete file {doc.file_path}: {e}")
                
                # 删除数据库记录
                await session.delete(doc)
                await session.commit()
                return True
            return False
    
    async def search_documents(self, category: str = None, tags: List[str] = None, status: str = None) -> List[Dict[str, Any]]:
        """搜索文档"""
        async for session in get_db_session():
            query = session.query(KnowledgeDocument)
            
            if category:
                query = query.filter(KnowledgeDocument.category == category)
            
            if tags:
                # 检查是否包含任意标签
                for tag in tags:
                    query = query.filter(KnowledgeDocument.tags.contains([tag]))
            
            if status:
                query = query.filter(KnowledgeDocument.status == status)
            
            documents = await query.all()
            
            return [
                {
                    "document_id": doc.document_id,
                    "filename": doc.filename,
                    "category": doc.category,
                    "tags": doc.tags,
                    "status": doc.status,
                    "created_at": doc.created_at,
                    "updated_at": doc.updated_at
                }
                for doc in documents
            ]
    
    async def get_document_categories(self) -> List[str]:
        """获取文档分类列表"""
        async for session in get_db_session():
            categories = await session.query(KnowledgeDocument.category).distinct().all()
            return [category[0] for category in categories if category[0]]
    
    async def get_document_tags(self) -> List[str]:
        """获取文档标签列表"""
        async for session in get_db_session():
            documents = await session.query(KnowledgeDocument.tags).all()
            tags = set()
            for doc_tags in documents:
                if doc_tags[0]:
                    tags.update(doc_tags[0])
            return sorted(list(tags))
    
    async def duplicate_document_check(self, file_hash: str) -> Optional[Dict[str, Any]]:
        """检查重复文档"""
        async for session in get_db_session():
            doc = await session.query(KnowledgeDocument).filter(KnowledgeDocument.file_hash == file_hash).first()
            if doc:
                return {
                    "document_id": doc.document_id,
                    "filename": doc.filename,
                    "created_at": doc.created_at
                }
            return None